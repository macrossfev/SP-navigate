"""Word document report exporter."""
from __future__ import annotations

import os
from typing import List, TYPE_CHECKING

from .base import BaseExporter
from .excel_exporter import _resolve_field

if TYPE_CHECKING:
    from navigate.core.models import PlanResult
    from navigate.core.config import NavigateConfig, ExportFormatConfig


class DocxExporter(BaseExporter):
    """Export plan results as a Word document report."""

    def export(self, result: "PlanResult", output_dir: str, **kwargs) -> str:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        os.makedirs(output_dir, exist_ok=True)
        tag = kwargs.get("tag", "")
        suffix = f"_{tag}" if tag else ""
        fmt_config: "ExportFormatConfig" = kwargs.get("format_config")
        title = (fmt_config.title if fmt_config and fmt_config.title
                 else "路线规划报告")
        path = os.path.join(output_dir, f"report{suffix}.docx")

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.page_width = Cm(29.7)
            section.page_height = Cm(21)

        style = doc.styles["Normal"]
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 1.3

        # Cover page
        for _ in range(4):
            doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run(title)
        run.font.size = Pt(28)
        run.bold = True

        doc.add_paragraph()
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"策略：{result.strategy_name}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()
        info_lines = [
            f"总点位：{result.total_points}",
            f"总天数：{result.total_days}",
            f"每日最大工时：{self.config.constraints.max_daily_hours} 小时",
            f"每点停留：{self.config.constraints.stop_time_per_point_min} 分钟",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)

        doc.add_page_break()

        # Summary table
        doc.add_heading("行程汇总", level=1)
        doc.add_paragraph()
        headers = ["天数", "点位", "距离 (km)", "驾驶 (min)", "总计 (h)"]
        st = doc.add_table(rows=len(result.days) + 2, cols=5,
                           style="Light Grid Accent 1")
        for j, h in enumerate(headers):
            c = st.cell(0, j)
            c.text = h
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for i, d in enumerate(result.days):
            row = st.rows[i + 1]
            row.cells[0].text = str(d.day)
            row.cells[1].text = str(d.point_count)
            row.cells[2].text = str(d.drive_distance_km)
            row.cells[3].text = str(d.drive_time_min)
            row.cells[4].text = str(d.total_time_hours)
            for c in row.cells:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.size = Pt(9)

        last = st.rows[-1]
        last.cells[0].text = "Total"
        last.cells[1].text = str(result.total_points)
        last.cells[2].text = str(round(result.total_distance_km, 1))
        last.cells[3].text = ""
        last.cells[4].text = str(round(result.total_hours, 1))
        for c in last.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        # Unassigned points
        if result.unassigned:
            doc.add_page_break()
            doc.add_heading("未分配点位（异常点）", level=1)
            p = doc.add_paragraph()
            threshold = self.config.strategy.options.get("outlier_threshold_km", 5.0)
            p.add_run(
                f"以下点位最近邻距离 > {threshold} km，未包含在主要行程中。"
            )
            doc.add_paragraph()
            ot = doc.add_table(rows=len(result.unassigned) + 1, cols=3,
                               style="Light Grid Accent 1")
            for j, h in enumerate(["#", "名称", "最近距离 (km)"]):
                c = ot.cell(0, j)
                c.text = h
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
            for i, (pt, nn_dist) in enumerate(result.unassigned):
                row = ot.rows[i + 1]
                row.cells[0].text = str(i + 1)
                row.cells[1].text = pt.name
                row.cells[2].text = f"{nn_dist:.1f}"
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)

        # Daily detail pages
        detail_fields = (fmt_config.detail_fields
                         if fmt_config and fmt_config.detail_fields else [])
        img_dir = os.path.join(output_dir, "images")

        for d in result.days:
            doc.add_page_break()
            trip_type_text = "隔夜住宿" if d.is_overnight else "单日往返"
            doc.add_heading(f"第{d.day}天 ({trip_type_text})", level=2)
            p = doc.add_paragraph()
            p.add_run(
                f"点位：{d.point_count} 个    "
                f"距离：{d.drive_distance_km} km    "
                f"总计：{d.total_time_hours} h"
            )

            # Include map image if exists
            if fmt_config and fmt_config.include_maps:
                for pattern in [f"day_{d.day}.png", f"Day{d.day}.png"]:
                    img_path = os.path.join(img_dir, pattern)
                    if os.path.exists(img_path):
                        ip = doc.add_paragraph()
                        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = ip.add_run()
                        run.add_picture(img_path, width=Cm(22))
                        break
                    else:
                        print(f"  Map image not found: {img_path}")

            doc.add_paragraph()

            if detail_fields:
                col_headers = [f.header for f in detail_fields]
                dt = doc.add_table(rows=len(d.points) + 1,
                                   cols=len(col_headers),
                                   style="Light Grid Accent 1")
                for j, h in enumerate(col_headers):
                    c = dt.cell(0, j)
                    c.text = h
                    for p in c.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(8)

                for idx, pt in enumerate(d.points):
                    row = dt.rows[idx + 1]
                    for j, f in enumerate(detail_fields):
                        val = _resolve_field(pt, d.day, idx + 1, f.source)
                        row.cells[j].text = str(val) if val else "-"
                        for c_p in row.cells[j].paragraphs:
                            for r in c_p.runs:
                                r.font.size = Pt(8)

        doc.save(path)
        print(f"[Export] Word: {path}")
        return path
