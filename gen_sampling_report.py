import os
import pandas as pd
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ========== 1. 读取数据 ==========

# 路线规划数据
with open("/root/projects/navigate/路线规划报告.json", "r", encoding="utf-8") as f:
    route_report = json.load(f)

# 最终地址列表（含坐标）
addr_df = pd.read_excel("/root/projects/navigate/最终地址列表.xlsx", engine="openpyxl")
addr_map = {}
for _, row in addr_df.iterrows():
    addr_map[str(row["地址"]).strip()] = {
        "matched": str(row["高德匹配"]).strip(),
        "coord": str(row["坐标"]).strip()
    }

# 二次供水现状表
survey_df = pd.read_excel("/root/projects/navigate/长寿区二次供水现状摸排统计2024.10.16(3).xls")

# 重命名列
survey_df.columns = [
    "序号", "小区名称", "泵房数量", "加压设施数量",
    "服务户数", "服务人口", "物业名称", "联系人", "联系电话",
    "管理方式", "水箱类型", "采样1", "采样2", "采样3"
]

# 跳过前两行（标题行和示例行），过滤有效数据
survey_df = survey_df.iloc[2:].reset_index(drop=True)
survey_df = survey_df[survey_df["小区名称"].notna() & (survey_df["小区名称"].str.strip() != "")]

# 构建小区信息字典（按小区名称匹配）
survey_map = {}
for _, row in survey_df.iterrows():
    name = str(row["小区名称"]).strip()
    survey_map[name] = {
        "序号": str(row["序号"]) if pd.notna(row["序号"]) else "",
        "泵房数量": str(row["泵房数量"]) if pd.notna(row["泵房数量"]) else "",
        "加压设施数量": str(row["加压设施数量"]) if pd.notna(row["加压设施数量"]) else "",
        "服务户数": str(row["服务户数"]) if pd.notna(row["服务户数"]) else "",
        "服务人口": str(row["服务人口"]) if pd.notna(row["服务人口"]) else "",
        "物业名称": str(row["物业名称"]) if pd.notna(row["物业名称"]) else "",
        "联系人": str(row["联系人"]) if pd.notna(row["联系人"]) else "",
        "联系电话": str(row["联系电话"]) if pd.notna(row["联系电话"]) else "",
        "管理方式": str(row["管理方式"]) if pd.notna(row["管理方式"]) else "",
        "水箱类型": str(row["水箱类型"]) if pd.notna(row["水箱类型"]) else "",
    }

# ========== 2. 修正地址 -> 原始名称 反查表 ==========

fix_df = pd.read_excel("/root/projects/navigate/地址修正表.xlsx", engine="openpyxl")
fix_col = fix_df.columns[-1]
addr_to_orig = {}
for _, row in fix_df.iterrows():
    orig = str(row["原始地址"]).replace("重庆市长寿区", "").strip()
    fix_addr = row[fix_col]
    if pd.notna(fix_addr) and str(fix_addr).strip():
        full_fix = str(fix_addr).strip()
        addr_to_orig[full_fix] = orig
        if not full_fix.startswith("重庆"):
            addr_to_orig[f"重庆市长寿区{full_fix}"] = orig
        else:
            addr_to_orig[full_fix] = orig

# ========== 3. 匹配函数 ==========

def find_survey_info(point_name):
    """根据导航点名称匹配二供现状表的信息，返回 (小区名称, info_dict) 或 (None, None)"""
    short = point_name.replace("重庆市长寿区", "").replace("重庆市", "").strip()

    # 1) 精确匹配
    if short in survey_map:
        return short, survey_map[short]

    # 2) 通过修正地址反查原始名称再匹配
    orig_name = addr_to_orig.get(point_name) or addr_to_orig.get(short)
    if orig_name and orig_name in survey_map:
        return orig_name, survey_map[orig_name]

    # 3) 标点符号归一化后匹配
    def normalize(s):
        return s.replace(".", "").replace("·", "").replace("·", "").replace("小区", "").replace("（", "(").replace("）", ")").strip()

    short_norm = normalize(short)
    for key, val in survey_map.items():
        key_norm = normalize(key)
        if key in short or short in key:
            return key, val
        if key_norm and short_norm and (key_norm in short_norm or short_norm in key_norm):
            return key, val

    return None, None

# ========== 3. 构建路线-采样点合并数据 ==========

day_details = []
unmatched = []

for day_info in route_report["days"]:
    day_num = day_info["day"]
    points = day_info["points"]  # 包含起点和返回点

    day_sampling_points = []
    for i, pt_name in enumerate(points):
        if i == 0 or i == len(points) - 1:
            continue  # 跳过起点和返回点

        community_name, survey = find_survey_info(pt_name)
        coord_info = addr_map.get(pt_name, {})

        entry = {
            "order": i,
            "name": pt_name,
            "short_name": pt_name.replace("重庆市长寿区", "").replace("重庆市", ""),
            "community_name": community_name or "",
            "coord": coord_info.get("coord", ""),
            "survey": survey
        }
        day_sampling_points.append(entry)

        if survey is None:
            unmatched.append(pt_name)

    day_details.append({
        "day": day_num,
        "point_count": day_info["point_count"],
        "drive_distance_km": day_info["drive_distance_km"],
        "drive_time_min": day_info["drive_time_min"],
        "stop_time_min": day_info["stop_time_min"],
        "total_time_hours": day_info["total_time_hours"],
        "sampling_points": day_sampling_points
    })

print(f"匹配完成。未匹配的点位: {len(unmatched)}")
for u in unmatched:
    print(f"  - {u}")

# ========== 4. 生成 Word 文档 ==========

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.page_width = Cm(29.7)  # A4横向
    section.page_height = Cm(21)

style = doc.styles['Normal']
style.font.size = Pt(10)
style.paragraph_format.line_spacing = 1.3

# ========== 封面 ==========
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("长寿区二次供水水质采样计划")
run.font.size = Pt(28)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("采样点 / 导航路线 / 联系人 综合报告")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    f"采样区域: 重庆市长寿区",
    f"总采样点: {route_report['total_points']} 个",
    f"总计天数: {route_report['total_days']} 天",
    f"起止地点: {route_report['base_point']}",
    f"每天限时: {route_report['max_daily_hours']} 小时",
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)

doc.add_page_break()

# ========== 一、基本信息 ==========
doc.add_heading("一、采样计划基本信息", level=1)
doc.add_paragraph()

info_table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
info_data = [
    ("采样区域", "重庆市长寿区（菩提街道、凤城街道、晏家街道、江南街道、渡舟街道、新市街道）"),
    ("起点/终点", route_report["base_point"]),
    ("采样点总数", f"{route_report['total_points']} 个"),
    ("计划天数", f"{route_report['total_days']} 天"),
    ("每天最大时间", f"{route_report['max_daily_hours']} 小时（含往返）"),
    ("每天最多点位", "5 个"),
    ("每点停留时间", f"{route_report['stop_time_per_point_min']} 分钟"),
]
for i, (k, v) in enumerate(info_data):
    info_table.cell(i, 0).text = k
    info_table.cell(i, 1).text = str(v)
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# ========== 二、日程汇总表 ==========
doc.add_heading("二、日程汇总表", level=1)
doc.add_paragraph()

sum_table = doc.add_table(rows=len(day_details) + 2, cols=5, style='Light Grid Accent 1')
sum_headers = ["天数", "采样点数", "驾车距离(km)", "驾车时间(min)", "总用时(小时)"]
for j, h in enumerate(sum_headers):
    cell = sum_table.cell(0, j)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

total_dist = 0
total_hours = 0
for i, dd in enumerate(day_details):
    row = sum_table.rows[i + 1]
    row.cells[0].text = f"第{dd['day']}天"
    row.cells[1].text = str(dd["point_count"])
    row.cells[2].text = str(dd["drive_distance_km"])
    row.cells[3].text = str(dd["drive_time_min"])
    row.cells[4].text = str(dd["total_time_hours"])
    total_dist += dd["drive_distance_km"]
    total_hours += dd["total_time_hours"]
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9)

last = sum_table.rows[-1]
last.cells[0].text = "合计"
last.cells[1].text = str(route_report["total_points"])
last.cells[2].text = str(round(total_dist, 1))
last.cells[3].text = ""
last.cells[4].text = str(round(total_hours, 1))
for cell in last.cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

doc.add_paragraph()

# ========== 三、每日采样详细计划 ==========
for dd in day_details:
    doc.add_page_break()
    doc.add_heading(f"第{dd['day']}天 采样计划", level=2)

    # 概览
    p = doc.add_paragraph()
    p.add_run(f"采样点数: {dd['point_count']}    "
              f"驾车距离: {dd['drive_distance_km']}km    "
              f"总用时: {dd['total_time_hours']}小时")

    # 插入导航轨迹图
    img_path = f"/root/projects/navigate/route_images/第{dd['day']}天导航轨迹图.png"
    if os.path.exists(img_path):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_p.add_run()
        run.add_picture(img_path, width=Cm(22))

    doc.add_paragraph()

    # 详细表格
    sp_list = dd["sampling_points"]
    num_rows = len(sp_list) + 1  # +1 for header

    col_headers = ["站序", "物业小区名称", "导航地址", "导航坐标", "物业名称", "联系人", "联系电话",
                   "服务户数", "服务人口", "加压设施", "管理方式"]
    num_cols = len(col_headers)

    detail_table = doc.add_table(rows=num_rows, cols=num_cols, style='Light Grid Accent 1')

    # 表头
    for j, h in enumerate(col_headers):
        cell = detail_table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

    # 数据行
    for i, sp in enumerate(sp_list):
        row = detail_table.rows[i + 1]
        survey = sp["survey"] or {}

        row.cells[0].text = f"第{sp['order']}站"
        row.cells[1].text = sp["community_name"] or sp["short_name"]
        row.cells[2].text = sp["short_name"] or sp["name"]
        row.cells[3].text = sp["coord"]
        row.cells[4].text = survey.get("物业名称", "-")
        row.cells[5].text = survey.get("联系人", "-")
        row.cells[6].text = survey.get("联系电话", "-")
        row.cells[7].text = survey.get("服务户数", "-")
        row.cells[8].text = survey.get("服务人口", "-")
        row.cells[9].text = survey.get("水箱类型", "-")
        row.cells[10].text = survey.get("管理方式", "-")

        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    # 设置列宽
    widths = [1.3, 3, 3.5, 3, 3.5, 1.3, 2.2, 1.3, 1.3, 2.2, 2.2]
    for row in detail_table.rows:
        for j, w in enumerate(widths):
            row.cells[j].width = Cm(w)

# ========== 四、全部采样点总表 ==========
doc.add_page_break()
doc.add_heading("四、全部采样点信息总表", level=1)
doc.add_paragraph()

# 收集所有采样点
all_points = []
for dd in day_details:
    for sp in dd["sampling_points"]:
        survey = sp["survey"] or {}
        all_points.append({
            "天数": f"第{dd['day']}天",
            "站序": f"第{sp['order']}站",
            "小区名称": sp["community_name"] or sp["short_name"],
            "导航地址": sp["short_name"],
            "坐标": sp["coord"],
            "物业": survey.get("物业名称", "-"),
            "联系人": survey.get("联系人", "-"),
            "电话": survey.get("联系电话", "-"),
            "户数": survey.get("服务户数", "-"),
            "人口": survey.get("服务人口", "-"),
            "设备": survey.get("水箱类型", "-"),
            "管理": survey.get("管理方式", "-"),
        })

total_headers = ["天数", "站序", "小区名称", "导航地址", "坐标", "物业", "联系人", "电话", "户数", "人口", "设备", "管理"]
total_table = doc.add_table(rows=len(all_points) + 1, cols=len(total_headers), style='Light Grid Accent 1')

for j, h in enumerate(total_headers):
    cell = total_table.cell(0, j)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(7)

for i, pt in enumerate(all_points):
    row = total_table.rows[i + 1]
    for j, key in enumerate(total_headers):
        row.cells[j].text = str(pt[key])
        for p in row.cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(7)

# ========== 保存 ==========
output_path = "/root/projects/navigate/采样计划报告.docx"
doc.save(output_path)
print(f"\n采样计划报告已保存: {output_path}")

# 同时输出 Excel 版总表
all_df = pd.DataFrame(all_points)
all_df.to_excel("/root/projects/navigate/采样计划总表.xlsx", index=False, engine="openpyxl")
print(f"采样计划总表已保存: /root/projects/navigate/采样计划总表.xlsx")
