"""
Word/Excel 报告生成模块
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .config import RouteConfig
from .strategies.base import PlanResult


class ReportGenerator:
    def __init__(self, config: RouteConfig):
        self.config = config
        os.makedirs(config.output_dir, exist_ok=True)

    def generate(self, result: PlanResult, tag: str = ""):
        """生成Word报告和Excel总表"""
        suffix = f"_{tag}" if tag else ""
        doc_path = os.path.join(self.config.output_dir, f"采样计划报告{suffix}.docx")
        xls_path = os.path.join(self.config.output_dir, f"采样计划总表{suffix}.xlsx")

        self._generate_word(result, doc_path, suffix)
        self._generate_excel(result, xls_path)

        print(f"[报告] Word: {doc_path}")
        print(f"[报告] Excel: {xls_path}")
        return doc_path, xls_path

    def _generate_word(self, result: PlanResult, path: str, suffix: str):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.page_width = Cm(29.7)
            section.page_height = Cm(21)

        style = doc.styles["Normal"]
        style.font.size = Pt(10)
        style.paragraph_format.line_spacing = 1.3

        # 封面
        for _ in range(4):
            doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = t.add_run("长寿区二次供水水质采样计划")
        run.font.size = Pt(28)
        run.bold = True

        doc.add_paragraph()
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(f"策略: {result.strategy_name}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()
        info_lines = [
            f"采样区域: 重庆市长寿区",
            f"总采样点: {result.total_points} 个",
            f"总计天数: {result.total_days} 天",
            f"每天限时: {result.config.max_daily_hours} 小时",
            f"每点停留: {result.config.stop_time_per_point_min} 分钟",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)

        doc.add_page_break()

        # 参数表
        doc.add_heading("一、规划参数", level=1)
        doc.add_paragraph()
        cfg = result.config
        param_data = [
            ("规划策略", result.strategy_name),
            ("起点/终点", cfg.base_point),
            ("每日最大时长", f"{cfg.max_daily_hours} 小时"),
            ("每点停留时间", f"{cfg.stop_time_per_point_min} 分钟"),
            ("往返基地开销", f"{cfg.roundtrip_overhead_min} 分钟"),
            ("缓冲系数", f"{cfg.buffer_factor}"),
            ("现场可用时间", f"{cfg.available_field_seconds/60:.0f} 分钟"),
            ("每天最多点位", f"{cfg.max_daily_points}"),
            ("平均车速", f"{cfg.avg_speed_kmh} km/h"),
            ("每日距离限制", "不限" if cfg.max_daily_distance_km == 0 else f"{cfg.max_daily_distance_km} km"),
        ]
        pt = doc.add_table(rows=len(param_data), cols=2, style="Light Grid Accent 1")
        for i, (k, v) in enumerate(param_data):
            pt.cell(i, 0).text = k
            pt.cell(i, 1).text = str(v)
            for cell in pt.rows[i].cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)

        doc.add_paragraph()

        # 日程汇总
        doc.add_heading("二、日程汇总表", level=1)
        doc.add_paragraph()
        headers = ["天数", "采样点数", "点间距离(km)", "驾车时间(min)", "总用时(小时)"]
        st = doc.add_table(rows=len(result.days) + 2, cols=5, style="Light Grid Accent 1")
        for j, h in enumerate(headers):
            c = st.cell(0, j)
            c.text = h
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for i, d in enumerate(result.days):
            row = st.rows[i + 1]
            row.cells[0].text = f"第{d.day}天"
            row.cells[1].text = str(len(d.point_indices))
            row.cells[2].text = str(d.drive_distance_km)
            row.cells[3].text = str(d.drive_time_min)
            row.cells[4].text = str(d.total_time_hours)
            for c in row.cells:
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.size = Pt(9)

        last = st.rows[-1]
        last.cells[0].text = "合计"
        last.cells[1].text = str(result.total_points)
        last.cells[2].text = str(round(result.total_distance_km, 1))
        last.cells[3].text = ""
        last.cells[4].text = str(round(result.total_hours, 1))
        for c in last.cells:
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        # 离群点说明
        if result.outliers:
            doc.add_page_break()
            doc.add_heading("三、离群点说明（未纳入采样计划）", level=1)
            p = doc.add_paragraph()
            threshold = cfg.outlier_threshold_km
            p.add_run(
                f"以下点位因地理位置孤立（最近邻距离超过 {threshold} km），"
                f"与其他采样点距离过远，单独前往效率极低，"
                f"建议另行安排专项采样或与周边区域联合采样。"
            )
            doc.add_paragraph()
            ot = doc.add_table(rows=len(result.outliers)+1, cols=4, style="Light Grid Accent 1")
            for j, h in enumerate(["序号", "点位名称", "导航地址", "最近邻距离(km)"]):
                c = ot.cell(0, j)
                c.text = h
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(9)
            for i, (sp, nn_dist) in enumerate(result.outliers):
                row = ot.rows[i+1]
                row.cells[0].text = str(i+1)
                row.cells[1].text = sp.community_name or sp.short_name
                row.cells[2].text = sp.short_name
                row.cells[3].text = f"{nn_dist:.1f}"
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(9)
            doc.add_paragraph()

        # 每日详细计划
        img_dir = os.path.join(self.config.output_dir, "images")
        for d in result.days:
            doc.add_page_break()
            doc.add_heading(f"第{d.day}天 采样计划", level=2)
            p = doc.add_paragraph()
            p.add_run(
                f"采样点数: {len(d.point_indices)}    "
                f"点间距离: {d.drive_distance_km}km    "
                f"总用时: {d.total_time_hours}小时"
            )

            # 自动检测点位图或导航轨迹图
            img_path = os.path.join(img_dir, f"第{d.day}天点位图.png")
            if not os.path.exists(img_path):
                img_path = os.path.join(img_dir, f"第{d.day}天导航轨迹图.png")
            if os.path.exists(img_path):
                ip = doc.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ip.add_run()
                run.add_picture(img_path, width=Cm(22))

            doc.add_paragraph()

            col_headers = ["站序", "小区名称", "导航地址", "物业名称",
                           "联系人", "联系电话", "服务户数", "服务人口", "管理方式"]
            dt = doc.add_table(rows=len(d.points) + 1, cols=len(col_headers),
                               style="Light Grid Accent 1")
            for j, h in enumerate(col_headers):
                c = dt.cell(0, j)
                c.text = h
                for p in c.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(8)

            for idx, sp in enumerate(d.points):
                row = dt.rows[idx + 1]
                row.cells[0].text = f"第{idx+1}站"
                row.cells[1].text = sp.community_name or sp.short_name
                row.cells[2].text = sp.short_name
                row.cells[3].text = sp.property_company or "-"
                row.cells[4].text = sp.contact_person or "-"
                row.cells[5].text = sp.contact_phone or "-"
                row.cells[6].text = sp.households or "-"
                row.cells[7].text = sp.population or "-"
                row.cells[8].text = sp.management or "-"
                for c in row.cells:
                    for p in c.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(8)

        doc.save(path)

    def _generate_excel(self, result: PlanResult, path: str):
        rows = []
        for d in result.days:
            for idx, sp in enumerate(d.points):
                rows.append({
                    "天数": f"第{d.day}天",
                    "站序": f"第{idx+1}站",
                    "小区名称": sp.community_name or sp.short_name,
                    "导航地址": sp.short_name,
                    "坐标": f"{sp.lng},{sp.lat}",
                    "物业": sp.property_company or "-",
                    "联系人": sp.contact_person or "-",
                    "电话": sp.contact_phone or "-",
                    "户数": sp.households or "-",
                    "人口": sp.population or "-",
                    "管理": sp.management or "-",
                    "当天距离km": d.drive_distance_km if idx == 0 else "",
                    "当天总时h": d.total_time_hours if idx == 0 else "",
                })
        pd.DataFrame(rows).to_excel(path, index=False, engine="openpyxl")
