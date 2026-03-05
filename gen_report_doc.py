import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 读取路线规划数据
with open("/root/projects/navigate/路线规划报告.json", "r", encoding="utf-8") as f:
    report = json.load(f)

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = 'SimSun'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("多点位路线规划报告")
run.font.size = Pt(26)
run.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("基于高德地图 API 自动规划")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"总点位: {report['total_points']} 个  |  总天数: {report['total_days']} 天")
run.font.size = Pt(12)

doc.add_page_break()

# ========== 目录页 ==========
h = doc.add_heading("目录", level=1)
doc.add_paragraph()

toc_items = [
    "一、基本信息",
    "二、路线规划汇总表",
]
for i in range(report["total_days"]):
    toc_items.append(f"三-{i+1}、第{i+1}天行程详情")

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ========== 一、基本信息 ==========
doc.add_heading("一、基本信息", level=1)
doc.add_paragraph()

info_table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ("起点/终点", report["base_point"]),
    ("总点位数", f"{report['total_points']} 个"),
    ("总天数", f"{report['total_days']} 天"),
    ("每天最大时间", f"{report['max_daily_hours']} 小时"),
    ("每天最多点位", "5 个"),
    ("每点停留时间", f"{report['stop_time_per_point_min']} 分钟"),
]

for i, (key, val) in enumerate(info_data):
    info_table.cell(i, 0).text = key
    info_table.cell(i, 1).text = str(val)
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ========== 二、路线规划汇总表 ==========
doc.add_heading("二、路线规划汇总表", level=1)
doc.add_paragraph()

summary_table = doc.add_table(rows=report["total_days"] + 2, cols=5, style='Light Grid Accent 1')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
headers = ["天数", "点位数", "驾车距离(km)", "驾车时间(min)", "总用时(小时)"]
for j, h in enumerate(headers):
    cell = summary_table.cell(0, j)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

total_dist = 0
total_time = 0

for i, day in enumerate(report["days"]):
    row = summary_table.rows[i + 1]
    row.cells[0].text = f"第{day['day']}天"
    row.cells[1].text = str(day["point_count"])
    row.cells[2].text = str(day["drive_distance_km"])
    row.cells[3].text = str(day["drive_time_min"])
    row.cells[4].text = str(day["total_time_hours"])
    total_dist += day["drive_distance_km"]
    total_time += day["total_time_hours"]
    for cell in row.cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 合计行
last_row = summary_table.rows[-1]
last_row.cells[0].text = "合计"
last_row.cells[1].text = str(report["total_points"])
last_row.cells[2].text = str(round(total_dist, 1))
last_row.cells[3].text = ""
last_row.cells[4].text = str(round(total_time, 1))
for cell in last_row.cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

doc.add_paragraph()

# ========== 三、每天行程详情 ==========
for day in report["days"]:
    doc.add_page_break()
    doc.add_heading(f"第{day['day']}天行程详情", level=2)

    # 概览
    p = doc.add_paragraph()
    p.add_run(f"点位数: ").bold = True
    p.add_run(f"{day['point_count']} 个    ")
    p.add_run(f"驾车距离: ").bold = True
    p.add_run(f"{day['drive_distance_km']} km    ")
    p.add_run(f"总用时: ").bold = True
    p.add_run(f"{day['total_time_hours']} 小时")

    p2 = doc.add_paragraph()
    p2.add_run(f"驾车时间: ").bold = True
    p2.add_run(f"{day['drive_time_min']} 分钟    ")
    p2.add_run(f"停留时间: ").bold = True
    p2.add_run(f"{day['stop_time_min']} 分钟")

    doc.add_paragraph()

    # 路线表格
    points = day["points"]
    detail_table = doc.add_table(rows=len(points), cols=3, style='Light Grid Accent 1')
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, pt_name in enumerate(points):
        if i == 0:
            role = "起点"
        elif i == len(points) - 1:
            role = "返回起点"
        else:
            role = f"第{i}站"

        detail_table.cell(i, 0).text = role
        detail_table.cell(i, 1).text = str(i)

        # 简化地址显示
        short_name = pt_name.replace("重庆市长寿区", "").replace("重庆市", "")
        if not short_name:
            short_name = pt_name
        detail_table.cell(i, 2).text = short_name

        for cell in detail_table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

    # 设置列宽
    for row in detail_table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(1.5)
        row.cells[2].width = Cm(12)

# ========== 保存 ==========
output_path = "/root/projects/navigate/路线规划报告.docx"
doc.save(output_path)
print(f"Word文档已保存: {output_path}")
